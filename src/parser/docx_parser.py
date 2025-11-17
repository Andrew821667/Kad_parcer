"""DOCX parser for court documents."""

import io
from typing import Any

from docx import Document

from src.core.exceptions import DOCXParseException
from src.core.logging import get_logger

logger = get_logger(__name__)


class DOCXDocumentParser:
    """Parser for DOCX court documents."""

    def __init__(self, docx_content: bytes) -> None:
        """Initialize parser with DOCX content.

        Args:
            docx_content: DOCX file content as bytes
        """
        self.docx_content = docx_content
        self.docx_file = io.BytesIO(docx_content)

        try:
            self.document = Document(self.docx_file)
        except Exception as e:
            raise DOCXParseException(f"Failed to load DOCX: {e}") from e

    def extract_text(self) -> str:
        """Extract text from DOCX.

        Returns:
            Extracted text

        Raises:
            DOCXParseException: If extraction fails
        """
        try:
            text_parts = []

            # Extract paragraphs
            for paragraph in self.document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            # Extract text from tables
            for table in self.document.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)

            full_text = "\n\n".join(text_parts)
            logger.debug(
                "docx_text_extracted",
                length=len(full_text),
                paragraphs=len(self.document.paragraphs),
                tables=len(self.document.tables),
            )

            return full_text

        except Exception as e:
            logger.error("docx_text_extraction_failed", error=str(e))
            raise DOCXParseException(f"Failed to extract text from DOCX: {e}") from e

    def _extract_table_text(self, table: Any) -> str:
        """Extract text from a table.

        Args:
            table: python-docx table object

        Returns:
            Extracted table text
        """
        rows = []

        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)

            if cells:
                rows.append(" | ".join(cells))

        return "\n".join(rows)

    def get_metadata(self) -> dict[str, Any]:
        """Extract DOCX metadata.

        Returns:
            Dictionary with metadata

        Raises:
            DOCXParseException: If extraction fails
        """
        try:
            metadata: dict[str, Any] = {
                "paragraph_count": len(self.document.paragraphs),
                "table_count": len(self.document.tables),
            }

            # Extract core properties
            core_props = self.document.core_properties
            metadata.update(
                {
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "comments": core_props.comments,
                    "created": core_props.created.isoformat() if core_props.created else None,
                    "modified": core_props.modified.isoformat() if core_props.modified else None,
                    "last_modified_by": core_props.last_modified_by,
                    "revision": core_props.revision,
                }
            )

            logger.debug("docx_metadata_extracted", metadata=metadata)
            return metadata

        except Exception as e:
            logger.error("docx_metadata_extraction_failed", error=str(e))
            raise DOCXParseException(f"Failed to extract DOCX metadata: {e}") from e

    def extract_paragraphs(self) -> list[str]:
        """Extract paragraphs as separate strings.

        Returns:
            List of paragraph texts

        Raises:
            DOCXParseException: If extraction fails
        """
        try:
            paragraphs = []

            for paragraph in self.document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            logger.debug("docx_paragraphs_extracted", count=len(paragraphs))
            return paragraphs

        except Exception as e:
            logger.error("docx_paragraphs_extraction_failed", error=str(e))
            raise DOCXParseException(f"Failed to extract paragraphs: {e}") from e
